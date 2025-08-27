from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.shared import OxmlElement, qn
from datetime import datetime, date, time, timedelta
import datetime
import streamlit as st
import ast
from io import BytesIO

def add_hyperlink(paragraph, text, url, font_name="Calibri", font_size=13.5, color=RGBColor(82, 119, 132)):
    """Ajoute un lien hypertexte à un paragraphe avec police, taille et couleur personnalisées"""
    
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )

    # Création de l'élément w:hyperlink
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Création du run et de ses propriétés
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Couleur
    c = OxmlElement('w:color')
    hex_color = "%02X%02X%02X" % (color[0], color[1], color[2])
    c.set(qn('w:val'), hex_color)
    rPr.append(c)

    # Souligné
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    # Police
    font = OxmlElement('w:rFonts')
    font.set(qn('w:ascii'), font_name)
    font.set(qn('w:hAnsi'), font_name)
    rPr.append(font)

    # Taille
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(font_size*2)))
    rPr.append(sz)

    new_run.append(rPr)

    # Ajouter le texte dans un w:t
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def format_participants(participants):
    """Formate la liste des participants"""
    if not participants:
        return "Participants non spécifiés"
    
    formatted = []
    for participant in participants:
        name = participant.get('name', 'Nom non spécifié')
        job = participant.get('job', '').strip()
        if job:
            formatted.append(f"{name}, {job}")
        else:
            formatted.append(f"{name}")
    
    return " ; ".join(formatted)

def get_french_weekday(date_obj):
    """Retourne le nom du jour en français"""
    days = {
        0: "Lundi",
        1: "Mardi", 
        2: "Mercredi",
        3: "Jeudi",
        4: "Vendredi",
        5: "Samedi",
        6: "Dimanche"
    }
    return days[date_obj.weekday()]

def get_french_month(date_obj):
    """Retourne le nom du mois en français"""
    months = {
        1: "janvier", 2: "février", 3: "mars", 4: "avril",
        5: "mai", 6: "juin", 7: "juillet", 8: "août", 
        9: "septembre", 10: "octobre", 11: "novembre", 12: "décembre"
    }
    return months[date_obj.month]

def format_french_date(date_obj):
    """Formate une date en français"""
    return f"{date_obj.day} {get_french_month(date_obj)}"

def get_date_range(events):
    """Calcule la plage de dates des événements (du lundi au vendredi de la semaine)"""
    if not events:
        return "Aucun événement"
    
    # Prendre n'importe quelle date des événements pour déterminer la semaine
    sample_date = events[0]['date']
    
    # Calculer le lundi de cette semaine (weekday() retourne 0 pour lundi)
    monday = sample_date - timedelta(days=sample_date.weekday())
    
    # Calculer le vendredi de cette semaine (lundi + 4 jours)
    friday = monday + timedelta(days=4)
    
    return f"Semaine du {format_french_date(monday)} au {format_french_date(friday)} {friday.year}"

def create_calendar_document(events_data, output_filename="calendrier_conferences.docx"):
    """Crée un document Word avec le calendrier des conférences"""
    
    doc = Document()
    
    # Configuration des marges
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Titre principal
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("CALENDRIER DES CONFÉRENCES".upper())
    title_run.bold = True
    title_run.font.size = Pt(17)
    title_run.font.color.rgb = RGBColor(79, 113, 190)
    title_run.font.name = 'Tahoma'
    
    # Sous-titre
    title_para.add_run().add_break(WD_BREAK.LINE)
    subtitle_run = title_para.add_run(get_date_range(events_data))
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.color.rgb = RGBColor(79, 113, 190)
    subtitle_run.font.name = 'Tahoma'
    
    # Grouper par date
    events_by_date = {}
    for event in events_data:
        event_date = event['date']
        events_by_date.setdefault(event_date, []).append(event)
    
    sorted_dates = sorted(events_by_date.keys())
    
    for event_date in sorted_dates:
        day_events = events_by_date[event_date]
        
        # En-tête du jour
        day_header = doc.add_paragraph()
        day_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        day_header_run = day_header.add_run(f"{get_french_weekday(event_date)} {format_french_date(event_date)}")
        day_header_run.bold = True
        day_header_run.font.size = Pt(15)
        day_header_run.font.name = 'Arial'
        day_header_run.font.color.rgb = RGBColor(150, 29, 19)
        
        for i, event in enumerate(day_events):
            # Titre + lien
            title_para = doc.add_paragraph()
            if event.get('link') and event.get('link').startswith('http'):
                add_hyperlink(title_para, event.get('title', 'Titre non spécifié'), event['link'], font_size=13.5)
                for run in title_para.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(13.5)
            else:
                title_run = title_para.add_run(event.get('title', 'Titre non spécifié'))
                title_run.font.name = 'Calibri'
                title_run.font.size = Pt(13.5)
            title_para.add_run().add_break()
            
            # Infos pratiques
            think_tank = event.get('think-tank', 'Think tank non spécifié')
            location = event.get('location', 'Localisation non spécifiée')
            hour = event.get('hour', 'Heure non spécifiée')
            info_text = f"{think_tank} — {location}. {hour.replace('-', '—')}"
            info_run = title_para.add_run(info_text)
            info_run.italic = True
            info_run.font.name = "Calibri"
            info_run.font.size = Pt(13.5)
            info_run.font.color.rgb = RGBColor(47, 110, 186)
            
            # Description justifiée
            description = event.get('description', 'Description non disponible')
            desc_para = doc.add_paragraph()
            desc_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            desc_run = desc_para.add_run(description)
            desc_run.font.name = 'Calibri'
            desc_run.font.size = Pt(13.5)
            
            # Participants en gras
            participants = event.get('participants', [])
            if participants:
                part_para = doc.add_paragraph()
                part_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                part_run = part_para.add_run("Participants : ")
                part_run.bold = True
                part_run.font.name = 'Calibri'
                part_run.font.size = Pt(13.5)	
                
                for j, participant in enumerate(participants):
                    name = participant.get('name', 'Nom non spécifié')
                    job = participant.get('job', '').strip()
                    
                    # Nom en gras
                    name_run = part_para.add_run(name)
                    name_run.bold = True
                    name_run.font.name = "Calibri"
                    name_run.font.size = Pt(13.5)
                    
                    # Job en texte normal
                    if job:
                        job_run = part_para.add_run(f", {job}")
                        job_run.font.name = "Calibri"
                        job_run.font.size = Pt(13.5)
                    
                    # Séparateur
                    if j < len(participants) - 1:
                        sep_run = part_para.add_run(" ; ")
                        sep_run.font.name = "Calibri"
                        sep_run.font.size = Pt(13.5)
            
            # Séparateur entre événements
            if i < len(day_events) - 1:
                separator = doc.add_paragraph("_" * 25)
                separator.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Séparateur entre jours
        if event_date != sorted_dates[-1]:
            doc.add_paragraph()
    
    return doc
    print(f"Document créé avec succès : {output_filename}")

st.title("📅 Formatteur Brain Juice")
user_input = st.text_area(
    "Collez votre liste de dictionnaires Python ci-dessous :",
    height=200,
    placeholder='Exemple : [{"date": datetime(2025, 9, 1), "event": "Rentrée"}]'
)

if st.button("Générer le document"):
    if not user_input.strip():
        st.warning("Veuillez entrer une liste de dictionnaires.")
    else:
        try:
            # Contexte sécurisé : seuls les noms autorisés seront évalués
            allowed_names = {
                "datetime": datetime,          # le module complet
                "date": datetime.date,         # la classe date
                "time": datetime.time,         # la classe time
            }

            data = eval(user_input, {"__builtins__": {}}, allowed_names)

            if not isinstance(data, list):
                st.error("L'entrée doit être une liste de dictionnaires.")
            else:
                # Création du document Word
                doc = create_calendar_document(data)

                # Sauvegarde en mémoire
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                # Téléchargement
                st.success("Document généré avec succès !")
                st.download_button(
                    label="📥 Télécharger le fichier Word",
                    data=buffer,
                    file_name="calendrier.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        except Exception as e:
            st.error(f"Erreur dans l'analyse des données : {e}")

















